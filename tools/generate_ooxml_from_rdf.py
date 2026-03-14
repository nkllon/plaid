from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt
from pyshacl import validate
from rdflib import Graph, Literal, Namespace, RDF


ROOT = Path("/Volumes/lemon/codex/plaid")
MODEL = ROOT / "model" / "security_package.ttl"
ONTOLOGY = ROOT / "model" / "security_ontology.ttl"
SHAPES = ROOT / "model" / "security_shapes.ttl"
OUTPUT_DIR = ROOT / "deliverables" / "docx"
TEXT_OUTPUT_DIR = ROOT / "deliverables" / "text"
SECDOC = Namespace("https://nkllon.com/plaid/security-doc#")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        if style_name in styles:
            styles[style_name].font.name = "Aptos"
    if "Title" in styles:
        styles["Title"].font.size = Pt(20)
    if "Heading 1" in styles:
        styles["Heading 1"].font.size = Pt(16)
    if "Heading 2" in styles:
        styles["Heading 2"].font.size = Pt(13)
    if "Heading 3" in styles:
        styles["Heading 3"].font.size = Pt(11.5)


def sort_key(graph: Graph, node, predicate) -> int:
    value = graph.value(node, predicate)
    return int(value) if isinstance(value, Literal) else 0


def text_value(graph: Graph, node, predicate=SECDOC.text) -> str:
    value = graph.value(node, predicate)
    return str(value) if value is not None else ""


def set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def set_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = min(section.page_width, section.page_height), max(
        section.page_width, section.page_height
    )
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def render_table(doc: Document, graph: Graph, table_node) -> None:
    columns = sorted(graph.objects(table_node, SECDOC.hasColumn), key=lambda n: sort_key(graph, n, SECDOC.columnOrder))
    rows = sorted(graph.objects(table_node, SECDOC.hasRow), key=lambda n: sort_key(graph, n, SECDOC.rowOrder))

    if not columns:
        return

    if len(columns) >= 6:
        set_landscape(doc.add_section(WD_SECTION.NEW_PAGE))

    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = True

    for c_idx, column in enumerate(columns):
        cell = table.cell(0, c_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.text = text_value(graph, column)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Aptos"
            run.font.size = Pt(8.5 if len(columns) >= 6 else 9)

    for r_idx, row in enumerate(rows, start=1):
        cells = sorted(graph.objects(row, SECDOC.hasCell), key=lambda n: sort_key(graph, n, SECDOC.cellOrder))
        for c_idx, cell_node in enumerate(cells[: len(columns)]):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = text_value(graph, cell_node)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Aptos"
                run.font.size = Pt(8 if len(columns) >= 6 else 9)

    doc.add_paragraph()

    if len(columns) >= 6:
        set_portrait(doc.add_section(WD_SECTION.CONTINUOUS))


def render_block(doc: Document, graph: Graph, block_node) -> None:
    block_types = set(graph.objects(block_node, RDF.type))
    if SECDOC.ParagraphBlock in block_types:
        doc.add_paragraph(text_value(graph, block_node))
        return

    if SECDOC.ListBlock in block_types:
        ordered = str(graph.value(block_node, SECDOC.isOrdered)).lower() == "true"
        items = sorted(graph.objects(block_node, SECDOC.hasItem), key=lambda n: sort_key(graph, n, SECDOC.itemOrder))
        style = "List Number" if ordered else "List Bullet"
        for item in items:
            doc.add_paragraph(text_value(graph, item), style=style)
        return

    if SECDOC.TableBlock in block_types:
        render_table(doc, graph, block_node)
        return


def render_document(doc: Document, graph: Graph, document_node, first: bool = False) -> None:
    if not first:
        doc.add_page_break()

    doc.add_paragraph(str(graph.value(document_node, SECDOC.title)), style="Title")
    sections = sorted(graph.objects(document_node, SECDOC.hasSection), key=lambda n: sort_key(graph, n, SECDOC.sectionOrder))

    for section_node in sections:
        heading = graph.value(section_node, SECDOC.title)
        level = int(graph.value(section_node, SECDOC.headingLevel))
        if heading:
            doc.add_paragraph(str(heading), style=f"Heading {min(level, 4)}")
        blocks = sorted(graph.objects(section_node, SECDOC.hasBlock), key=lambda n: sort_key(graph, n, SECDOC.blockOrder))
        for block in blocks:
            render_block(doc, graph, block)


def render_block_text(graph: Graph, block_node) -> list[str]:
    block_types = set(graph.objects(block_node, RDF.type))
    if SECDOC.ParagraphBlock in block_types:
        return [text_value(graph, block_node), ""]

    if SECDOC.ListBlock in block_types:
        ordered = str(graph.value(block_node, SECDOC.isOrdered)).lower() == "true"
        items = sorted(graph.objects(block_node, SECDOC.hasItem), key=lambda n: sort_key(graph, n, SECDOC.itemOrder))
        lines = []
        for idx, item in enumerate(items, start=1):
            prefix = f"{idx}. " if ordered else "- "
            lines.append(prefix + text_value(graph, item))
        lines.append("")
        return lines

    if SECDOC.TableBlock in block_types:
        columns = sorted(graph.objects(block_node, SECDOC.hasColumn), key=lambda n: sort_key(graph, n, SECDOC.columnOrder))
        rows = sorted(graph.objects(block_node, SECDOC.hasRow), key=lambda n: sort_key(graph, n, SECDOC.rowOrder))
        header = " | ".join(text_value(graph, col) for col in columns)
        separator = " | ".join("---" for _ in columns)
        lines = [header, separator]
        for row in rows:
            cells = sorted(graph.objects(row, SECDOC.hasCell), key=lambda n: sort_key(graph, n, SECDOC.cellOrder))
            lines.append(" | ".join(text_value(graph, cell) for cell in cells))
        lines.append("")
        return lines

    return []


def render_document_text(graph: Graph, document_node) -> str:
    lines = [str(graph.value(document_node, SECDOC.title)), ""]
    sections = sorted(graph.objects(document_node, SECDOC.hasSection), key=lambda n: sort_key(graph, n, SECDOC.sectionOrder))
    for section_node in sections:
        heading = graph.value(section_node, SECDOC.title)
        if heading:
            lines.append(str(heading))
            lines.append("")
        blocks = sorted(graph.objects(section_node, SECDOC.hasBlock), key=lambda n: sort_key(graph, n, SECDOC.blockOrder))
        for block in blocks:
            lines.extend(render_block_text(graph, block))
    return "\n".join(lines).strip() + "\n"


def load_graph() -> Graph:
    data_graph = Graph()
    for source in (ONTOLOGY, MODEL):
        data_graph.parse(source)

    shapes_graph = Graph()
    shapes_graph.parse(SHAPES)

    conforms, _, report = validate(
        data_graph,
        shacl_graph=shapes_graph,
        ont_graph=None,
        inference="rdfs",
        abort_on_first=False,
        allow_infos=False,
        allow_warnings=False,
    )
    if not conforms:
        raise SystemExit(str(report))
    return data_graph


def build_single_docs(graph: Graph) -> None:
    documents = sorted(graph.subjects(RDF.type, SECDOC.Document), key=lambda n: sort_key(graph, n, SECDOC.documentOrder))
    for document_node in documents:
        file_stem = str(graph.value(document_node, SECDOC.fileStem))
        out = OUTPUT_DIR / f"{file_stem}.docx"
        doc = Document()
        configure_document(doc)
        render_document(doc, graph, document_node, first=True)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out)


def build_combined_doc(graph: Graph) -> None:
    documents = sorted(graph.subjects(RDF.type, SECDOC.Document), key=lambda n: sort_key(graph, n, SECDOC.documentOrder))
    doc = Document()
    configure_document(doc)
    for index, document_node in enumerate(documents):
        render_document(doc, graph, document_node, first=index == 0)
    out = OUTPUT_DIR / "Plaid_Security_Documentation_Package.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def build_text_docs(graph: Graph) -> None:
    documents = sorted(graph.subjects(RDF.type, SECDOC.Document), key=lambda n: sort_key(graph, n, SECDOC.documentOrder))
    TEXT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for document_node in documents:
        file_stem = str(graph.value(document_node, SECDOC.fileStem))
        out = TEXT_OUTPUT_DIR / f"{file_stem}.txt"
        out.write_text(render_document_text(graph, document_node))


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    graph = load_graph()
    build_single_docs(graph)
    build_combined_doc(graph)
    build_text_docs(graph)


if __name__ == "__main__":
    main()
